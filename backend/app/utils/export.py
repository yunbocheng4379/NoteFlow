import os
import re
from dotenv import load_dotenv

load_dotenv()

# 项目根路径（无论你在哪里运行）
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 从 .env 获取 DATA_DIR，相对于 BASE_DIR 解析
DATA_DIR_NAME = os.getenv("DATA_DIR", "data")
DATA_DIR = os.path.join(BASE_DIR, DATA_DIR_NAME)
SAVE_PATH = os.path.join(DATA_DIR, "note_output")
IMAGE_BASE_URL = os.getenv("IMAGE_BASE_URL")
STATIC_BASE = os.path.join(BASE_DIR, IMAGE_BASE_URL) if IMAGE_BASE_URL else BASE_DIR

# 内容时间戳标记，例如 *Content-[04:16]* / Screenshot-04:16（正常流程里会在生成笔记时
# 被替换为跳转链接/截图，这里兜底处理未被替换的残留标记，渲染成醒目的徽章而非裸文字）
_TS_MARKER_PATTERN = re.compile(
    r"\*?(?:Content|Screenshot)-\[?(\d{2}):(\d{2})\]?\*?"
)

# 统一导出样式：h1/h2/h3 层级色差、代码块、引用块、表格、时间戳徽章、原片跳转标签
_PAGE_STYLE = """
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
       max-width: 860px; margin: 40px auto; padding: 0 24px;
       color: #1a1a1a; line-height: 1.75; }}
h1, h2, h3, h4, h5, h6 {{ font-weight: 700; line-height: 1.35; }}
h1 {{ font-size: 1.9em; color: #111827; margin: 0 0 .6em; padding-bottom: .5em;
     border-bottom: 3px solid #111827; }}
h2 {{ font-size: 1.45em; color: #1f2937; margin: 1.6em 0 .6em; padding-bottom: .35em;
     border-bottom: 1px solid #e5e7eb; }}
h3 {{ font-size: 1.2em; color: #374151; margin: 1.3em 0 .5em; }}
h4, h5, h6 {{ font-size: 1.05em; color: #4b5563; margin: 1.1em 0 .4em; }}
p {{ margin: .8em 0; }}
.codehilite, pre {{ background: #282c34; color: #abb2bf; border-radius: 8px;
     padding: 16px; overflow-x: auto; margin: 1em 0; }}
.codehilite pre {{ margin: 0; background: none; padding: 0; }}
code {{ background: #f0f1f3; border-radius: 4px; padding: 2px 6px; font-size: .9em;
     color: #c7254e; }}
pre code, .codehilite code {{ background: none; padding: 0; color: inherit; }}
blockquote {{ border-left: 4px solid #6366f1; margin: 1em 0; padding: .4em 1em;
     background: #f5f5ff; color: #4b5563; border-radius: 0 6px 6px 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 1.2em 0; font-size: .95em; }}
th, td {{ border: 1px solid #e5e7eb; padding: 8px 14px; text-align: left; }}
th {{ background: #f3f4f6; font-weight: 600; }}
tr:nth-child(even) td {{ background: #fafafa; }}
img {{ max-width: 100%; border-radius: 8px; }}
a {{ color: #0f766e; text-decoration: none; }}
hr {{ border: none; border-top: 1px solid #e5e7eb; margin: 2em 0; }}
.ts-badge {{ display: inline-block; background: #eef2ff; color: #4338ca;
     border-radius: 999px; padding: 1px 10px; font-size: .85em; font-weight: 600; }}
.origin-link {{ display: inline-block; background: #ecfdf5; color: #047857;
     border-radius: 999px; padding: 1px 10px; font-size: .85em; font-weight: 600; }}
{pygments_css}
"""


class ExportUtils:
    def __init__(self, **kwargs):
        print(f"保存路径: {SAVE_PATH}")
        print(f"静态文件路径: {STATIC_BASE}")
        if not os.path.exists(SAVE_PATH):
            os.makedirs(SAVE_PATH)

    def _embed_image_as_base64(self, img_path: str) -> str:
        import base64
        import mimetypes

        try:
            mime_type, _ = mimetypes.guess_type(img_path)
            if not mime_type:
                ext = os.path.splitext(img_path)[1].lower()
                mime_map = {
                    '.png': 'image/png',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.gif': 'image/gif',
                    '.bmp': 'image/bmp',
                    '.webp': 'image/webp',
                    '.svg': 'image/svg+xml'
                }
                mime_type = mime_map.get(ext, 'image/png')

            with open(img_path, 'rb') as f:
                img_data = f.read()

            base64_data = base64.b64encode(img_data).decode('utf-8')
            return f"data:{mime_type};base64,{base64_data}"

        except Exception as e:
            print(f"图片 base64 编码失败 {img_path}: {str(e)}")
            return None

    def _get_normalized_path(self, path: str) -> str:
        return os.path.normpath(os.path.abspath(path))

    def _replace_static_paths_with_absolute(self, content: str) -> str:
        def repl(match):
            alt_text = match.group(1) if match.group(1) else ""
            img_path = match.group(2).strip()

            print(f"处理图片路径: {img_path}")

            if img_path.startswith("/static/"):
                relative_path = img_path.lstrip("/")
                abs_path = os.path.join(BASE_DIR, relative_path)
                abs_path = self._get_normalized_path(abs_path)

                if os.path.exists(abs_path):
                    base64_uri = self._embed_image_as_base64(abs_path)
                    if base64_uri:
                        return f"![{alt_text}]({base64_uri})"
                    return f"![{alt_text}](图片转换失败: {img_path})"
                return f"![{alt_text}](图片不存在: {img_path})"

            elif not img_path.startswith(('http://', 'https://', 'data:')):
                possible_paths = [
                    os.path.join(STATIC_BASE, img_path),
                    os.path.abspath(img_path),
                    os.path.join(BASE_DIR, img_path)
                ]

                for abs_path in possible_paths:
                    abs_path = self._get_normalized_path(abs_path)
                    if os.path.exists(abs_path):
                        base64_uri = self._embed_image_as_base64(abs_path)
                        if base64_uri:
                            return f"![{alt_text}]({base64_uri})"
                        break

                return f"![{alt_text}](图片未找到: {img_path})"

            return match.group(0)

        pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        return re.sub(pattern, repl, content)

    def _render_markers(self, html_body: str) -> str:
        """把残留的时间戳标记转成徽章样式，原片跳转链接转成绿色小标签。"""
        html_body = _TS_MARKER_PATTERN.sub(
            lambda m: f'<span class="ts-badge">▶ {m.group(1)}:{m.group(2)}</span>',
            html_body,
        )
        html_body = re.sub(
            r'(<a\s+[^>]*>)(原片 @ [^<]+)(</a>)',
            lambda m: f'{m.group(1)[:-1]} class="origin-link">{m.group(2)}{m.group(3)}',
            html_body,
        )
        return html_body

    def _render_html_body(self, content: str) -> str:
        import markdown

        html_body = markdown.markdown(
            content,
            extensions=["extra", "toc", "tables", "fenced_code", "codehilite"],
            extension_configs={"codehilite": {"guess_lang": False}},
        )
        return self._render_markers(html_body)

    def _pygments_css(self) -> str:
        from pygments.formatters import HtmlFormatter

        return HtmlFormatter(style="monokai").get_style_defs(".codehilite")

    def _full_html(self, content: str, title: str) -> str:
        html_body = self._render_html_body(content)
        page_style = _PAGE_STYLE.format(pygments_css=self._pygments_css())

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{title}</title>
  <style>
{page_style}
  </style>
</head>
<body>
{html_body}
</body>
</html>"""

    def _to_html(self, content: str, title: str) -> str:
        full_html = self._full_html(content, title)
        save_path = os.path.join(SAVE_PATH, f"{title}.html")
        with open(save_path, "w", encoding="utf-8") as f:
            f.write(full_html)
        return save_path

    def _find_chrome(self) -> str | None:
        import shutil

        candidates = [
            os.environ.get("CHROME_BIN_PATH"),
            "google-chrome", "google-chrome-stable", "chromium", "chromium-browser",
            "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
            "/Applications/Chromium.app/Contents/MacOS/Chromium",
            "/usr/bin/google-chrome", "/usr/bin/chromium", "/usr/bin/chromium-browser",
        ]
        for candidate in candidates:
            if not candidate:
                continue
            if os.path.isabs(candidate):
                if os.path.exists(candidate):
                    return candidate
                continue
            found = shutil.which(candidate)
            if found:
                return found
        return None

    def _to_pdf(self, content: str, title: str) -> str:
        """Markdown → HTML → Chrome headless print-to-pdf，与 _to_image 共享同一套渲染 CSS。"""
        import subprocess

        chrome = self._find_chrome()
        save_path = os.path.join(SAVE_PATH, f"{title}.pdf")
        if not chrome:
            raise RuntimeError(
                "PDF 导出需要 Chrome 或 Chromium，未在系统中找到。"
                "请安装 Google Chrome 后重试。"
            )

        html_path = self._to_html(content, f"_tmp_{title}")
        try:
            result = subprocess.run(
                [
                    chrome,
                    "--headless=new",
                    "--disable-gpu",
                    "--no-sandbox",
                    "--disable-dev-shm-usage",
                    f"--print-to-pdf={save_path}",
                    "--no-pdf-header-footer",
                    f"file://{os.path.abspath(html_path)}",
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0 or not os.path.exists(save_path):
                raise RuntimeError(f"Chrome PDF 导出失败：{result.stderr[:300]}")
        finally:
            try:
                os.remove(html_path)
            except Exception:
                pass

        return save_path

    # ---- Word 导出：基于 _render_html_body() 产出的 HTML DOM 渲染，而非逐行正则 ----

    def _set_paragraph_shading(self, paragraph, hex_color: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        paragraph._p.get_or_add_pPr().append(shd)

    def _set_paragraph_border_bottom(self, paragraph) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "D1D5DB")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _set_cell_shading(self, cell, hex_color: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)

    def _add_image_to_paragraph(self, paragraph, src: str) -> None:
        import base64
        import io
        import urllib.request
        from docx.shared import Inches

        try:
            if src.startswith("data:"):
                _, b64data = src.split(",", 1)
                stream = io.BytesIO(base64.b64decode(b64data))
            elif src.startswith(("http://", "https://")):
                with urllib.request.urlopen(src, timeout=10) as resp:
                    stream = io.BytesIO(resp.read())
            else:
                return
            paragraph.add_run().add_picture(stream, width=Inches(5.5))
        except Exception as e:
            print(f"Word 图片插入失败: {e}")

    def _add_word_inline(self, paragraph, node, bold: bool = False, italic: bool = False) -> None:
        from bs4 import NavigableString
        from docx.shared import RGBColor

        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
            return

        name = node.name
        classes = node.get("class") or [] if hasattr(node, "get") else []

        if name in ("strong", "b"):
            for child in node.children:
                self._add_word_inline(paragraph, child, bold=True, italic=italic)
        elif name in ("em", "i"):
            for child in node.children:
                self._add_word_inline(paragraph, child, bold=bold, italic=True)
        elif name == "code":
            run = paragraph.add_run(node.get_text())
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            run.bold = bold
            run.italic = italic
        elif name == "span" and "ts-badge" in classes:
            run = paragraph.add_run(node.get_text())
            run.font.color.rgb = RGBColor(0x43, 0x38, 0xCA)
            run.bold = True
        elif name == "span" and "origin-link" in classes:
            run = paragraph.add_run(node.get_text())
            run.font.color.rgb = RGBColor(0x04, 0x78, 0x57)
            run.bold = True
        elif name == "a":
            run = paragraph.add_run(node.get_text())
            run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            run.underline = True
            run.bold = bold
            run.italic = italic
        elif name == "img":
            self._add_image_to_paragraph(paragraph, node.get("src", ""))
        elif name == "br":
            paragraph.add_run().add_break()
        else:
            for child in getattr(node, "children", []):
                self._add_word_inline(paragraph, child, bold=bold, italic=italic)

    def _render_word_block(self, doc, node) -> None:
        from docx.shared import Pt, RGBColor, Inches
        from bs4 import NavigableString

        if isinstance(node, NavigableString) or not getattr(node, "name", None):
            return

        name = node.name
        HEADING_COLORS = {
            1: RGBColor(0x11, 0x18, 0x27), 2: RGBColor(0x1F, 0x29, 0x37),
            3: RGBColor(0x37, 0x41, 0x51), 4: RGBColor(0x4B, 0x55, 0x63),
        }
        HEADING_SIZES = {1: Pt(22), 2: Pt(17), 3: Pt(14), 4: Pt(12)}

        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = min(int(name[1]), 4)
            heading = doc.add_heading("", level=level)
            for child in node.children:
                self._add_word_inline(heading, child, bold=True)
            for run in heading.runs:
                run.font.color.rgb = HEADING_COLORS[level]
                run.font.size = HEADING_SIZES[level]
        elif name == "p":
            para = doc.add_paragraph()
            for child in node.children:
                self._add_word_inline(para, child)
        elif name == "pre":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.2)
            run = para.add_run(node.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            self._set_paragraph_shading(para, "F0F1F3")
        elif name == "blockquote":
            children = [c for c in node.children if getattr(c, "name", None) or str(c).strip()]
            for child in children:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.4)
                if getattr(child, "name", None):
                    for grandchild in child.children:
                        self._add_word_inline(para, grandchild, italic=True)
                else:
                    self._add_word_inline(para, child, italic=True)
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        elif name in ("ul", "ol"):
            style_name = "List Bullet" if name == "ul" else "List Number"
            for li in node.find_all("li", recursive=False):
                para = doc.add_paragraph(style=style_name)
                for child in li.children:
                    if getattr(child, "name", None) in ("ul", "ol"):
                        continue
                    self._add_word_inline(para, child)
                for nested in li.find_all(["ul", "ol"], recursive=False):
                    self._render_word_block(doc, nested)
        elif name == "table":
            rows = node.find_all("tr")
            if not rows:
                return
            n_cols = max(len(r.find_all(["td", "th"])) for r in rows)
            table = doc.add_table(rows=0, cols=n_cols)
            table.style = "Table Grid"
            for r_idx, tr in enumerate(rows):
                cells = tr.find_all(["td", "th"])
                row_cells = table.add_row().cells
                is_header = r_idx == 0 and all(c.name == "th" for c in cells)
                for c_idx, cell in enumerate(cells):
                    if c_idx >= n_cols:
                        break
                    para = row_cells[c_idx].paragraphs[0]
                    for child in cell.children:
                        self._add_word_inline(para, child, bold=is_header)
                    if is_header:
                        self._set_cell_shading(row_cells[c_idx], "F3F4F6")
                    elif r_idx % 2 == 0:
                        self._set_cell_shading(row_cells[c_idx], "FAFAFA")
        elif name == "hr":
            self._set_paragraph_border_bottom(doc.add_paragraph())
        elif name == "img":
            self._add_image_to_paragraph(doc.add_paragraph(), node.get("src", ""))
        else:
            for child in getattr(node, "children", []):
                self._render_word_block(doc, child)

    def _to_word(self, content: str, title: str) -> str:
        from docx import Document
        from docx.shared import Pt
        from bs4 import BeautifulSoup

        html_body = self._render_html_body(content)
        soup = BeautifulSoup(html_body, "html.parser")

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style.font.size = Pt(11)

        for node in soup.children:
            self._render_word_block(doc, node)

        save_path = os.path.join(SAVE_PATH, f"{title}.docx")
        doc.save(save_path)
        return save_path

    def _to_image(self, content: str, title: str) -> str:
        """Markdown → HTML → Chrome headless 截图 → PNG"""
        import subprocess

        chrome = self._find_chrome()
        if not chrome:
            raise RuntimeError(
                "图片导出需要 Chrome 或 Chromium，未在系统中找到。"
                "请安装 Google Chrome 后重试。"
            )

        html_path = self._to_html(content, f"_tmp_{title}")
        save_path = os.path.join(SAVE_PATH, f"{title}.png")

        try:
            result = subprocess.run(
                [
                    chrome,
                    "--headless=new",
                    "--disable-gpu",
                    "--no-sandbox",
                    "--disable-dev-shm-usage",
                    "--hide-scrollbars",
                    "--window-size=1200,800",
                    f"--screenshot={save_path}",
                    f"file://{os.path.abspath(html_path)}",
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0 or not os.path.exists(save_path):
                raise RuntimeError(f"Chrome 截图失败：{result.stderr[:300]}")
        finally:
            try:
                os.remove(html_path)
            except Exception:
                pass

        return save_path

    def export(self, output_format: str, title: str, content: str) -> str:
        content = content.strip()
        print("开始处理图片路径...")
        content = self._replace_static_paths_with_absolute(content)
        output_format = output_format.lower()

        if output_format == "pdf":
            return self._to_pdf(content, title)
        elif output_format == "html":
            return self._to_html(content, title)
        elif output_format in ["word", "docx"]:
            return self._to_word(content, title)
        elif output_format in ["image", "png"]:
            return self._to_image(content, title)
        elif output_format == "md":
            save_path = os.path.join(SAVE_PATH, f"{title}.md")
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(content)
            return save_path
        else:
            raise ValueError(f"不支持的导出格式: {output_format}")

    def get_supported_formats(self):
        return {
            "md": "Markdown 文档",
            "pdf": "PDF 文档",
            "html": "HTML 网页",
            "docx": "Word 文档 (.docx)",
            "png": "PNG 图片",
        }
